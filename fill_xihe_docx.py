from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn


SOURCE = "/Users/tyrion/Documents/员工关怀-生日提醒/附件5：应用创新大赛项目创意书-羲和完善版.docx"


SECTIONS = {
    1: [
        ("团队定位", "项目采用“小核心、快迭代”的Vibe Coding协作模式，由王一喆担任项目负责人，统筹产品定位、需求分析、技术架构和成果汇报。"),
        ("职责分工", "负责人完成业务调研、原型设计、AI辅助编码、测试验证及迭代管理；后续根据试点需要，引入员工关怀业务、交互设计、数据安全和运营推广人员共同完善。"),
        ("协作机制", "以真实用户任务为牵引，按照“需求澄清—视觉验证—交互原型—小范围试用—持续改进”推进，确保产品价值、体验与技术可行性同步验证。"),
    ],
    2: [
        ("业务场景", "亲人、朋友和同事的生日信息分散在通讯录、聊天记录和表格中，用户通常依赖记忆或临时提醒；企业员工关怀也存在名单维护分散、提醒不及时、表达方式单一等问题。"),
        ("主要痛点", "手工录入成本高，公历与农历并存容易出错；普通日历只提醒“当天”，缺少礼物、聚餐和祝福的准备时间；提醒后还需在多个应用间切换，难以形成关怀闭环。"),
        ("产品机会", "“羲和”以日光时序为品牌意象，面向移动端提供温暖、克制、可信赖的重要日子管理体验。"),
    ],
    3: [
        ("一是数据难沉淀", "通过手动录入、CSV/XLSX导入、字段映射、重复识别和自定义分组，降低生日信息整理成本。"),
        ("二是提醒不可靠", "建立“提前3天+当天”的默认提醒节奏，兼容公历、农历和自定义时间，并预留微信等多渠道触达能力。"),
        ("三是关怀难落地", "提醒卡片直接提供祝福、联系、礼物待办和“已关心”反馈，使用户从记住日期自然过渡到完成表达。"),
        ("四是数据缺乏信任", "坚持最少采集、权限透明、可导出可删除，让个人关系数据始终处于用户控制之下。"),
    ],
    4: [
        ("产品形态", "采用移动端优先的响应式Web/PWA方案，首阶段以高保真可交互原型验证首页、未来三日、人物分组、文件导入、提醒设置和祝福闭环。"),
        ("系统架构", "前端负责交互与本地缓存；后端拟提供账号、人物、分组、生日、提醒任务和操作记录等服务；生日计算引擎统一处理公历、农历、时区与周期规则。"),
        ("服务集成", "文件导入模块完成解析、映射、校验和去重；消息适配层对接系统通知，并按平台规则扩展微信服务通知、短信或邮件。"),
        ("安全方案", "采用传输加密、最小权限、敏感字段保护、备份恢复和操作留痕。"),
    ],
    5: [
        ("Vibe Coding", "利用大模型辅助完成需求拆解、界面生成、代码实现和测试修正，由人工负责目标约束、架构决策与质量验收，缩短从创意到原型的周期。"),
        ("生日规则引擎", "处理公历/农历换算、闰月策略、未知年份、跨年窗口及时区变化，保证重复提醒计算准确。"),
        ("任务与推送", "采用定时任务、幂等控制、失败重试和通知健康检查，降低漏提醒、重复提醒风险；通过统一适配层支持多渠道扩展。"),
        ("体验技术", "使用React响应式组件与轻量动效实现移动端交互，结合结构化数据校验和本地缓存提升可用性。"),
    ],
    6: [
        ("理念创新", "由“生日提醒器”升级为“轻量关系关怀助手”，形成“提前发现—做好准备—自然表达—留下记录”的完整闭环。"),
        ("机制创新", "以未来三天作为核心准备窗口，将人物、剩余时间、祝福、联系和礼物待办集中在同一提醒场景，减少应用跳转和行动断层。"),
        ("体验创新", "融入“羲和”日光时序意象，以东方极简、时间刻度和柔和光影建立区别于传统日历工具的品牌体验。"),
        ("研发创新", "通过Vibe Coding快速完成产品规划、视觉方向和移动端交互原型，为小团队低成本验证AI应用提供可复制路径。"),
    ],
    7: [
        ("业务成效", "帮助个人减少重要生日遗漏，提前完成祝福和礼物准备；在企业场景中，可为员工生日关怀提供统一数据、主动提醒和标准化执行入口，提升组织温度。"),
        ("产品成果", "形成可在手机端演示的交互原型，并逐步交付支持文件导入、分组、农历生日、提醒和数据备份的MVP；以“60秒录入首位人物、10秒完成提醒后行动”为体验目标。"),
        ("商业潜力", "个人版可探索增值订阅，企业版可扩展团队生日册、分级权限、祝福模板和员工关怀运营，具备从内部工具向标准化SaaS产品演进的可能。"),
    ],
}


def set_run_font(run, size=10.5, bold=False):
    run.font.name = "宋体"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def clear_cell(cell):
    for paragraph in cell.paragraphs[1:]:
        paragraph._element.getparent().remove(paragraph._element)
    paragraph = cell.paragraphs[0]
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)
    return paragraph


def fill_response_cell(cell, items):
    first = clear_cell(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for index, (label, body) in enumerate(items):
        paragraph = first if index == 0 else cell.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(2 if index else 0)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        paragraph.paragraph_format.line_spacing = 1.15
        label_run = paragraph.add_run(f"{label}：")
        set_run_font(label_run, bold=True)
        body_run = paragraph.add_run(body)
        set_run_font(body_run)


doc = Document(SOURCE)
for table_index, items in SECTIONS.items():
    fill_response_cell(doc.tables[table_index].cell(1, 0), items)

# Preserve the template and existing team data; only refresh editable metadata.
doc.core_properties.title = "羲和（智能生日管家）—应用创新大赛项目创意书"
doc.core_properties.subject = "基于Vibe Coding的移动端智能生日与关系关怀产品"
doc.save(SOURCE)
print(SOURCE)
